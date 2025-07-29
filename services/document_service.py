# services/document_service.py
import io
import base64
import re
from typing import Dict, Any
from fastapi import HTTPException
import openpyxl
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from models.schemas import BaseAIRequest, FinalDocumentsResponse
from services.ai_service import ai_service

class DocumentService:
    """Service for generating research documents"""
    
    def __init__(self):
        pass
    
    def _process_text_for_word(self, text: str, paragraph) -> None:
        """
        Process text with LaTeX-style formatting and add properly formatted runs to a Word paragraph.
        Handles subscripts, superscripts, bold, italic, and chemical formulas.
        """
        # Split text by formatting patterns while preserving the delimiters
        pattern = r'(\$_\{[^}]+\}\$|\$_[^$\s]+\$|\$\^\{[^}]+\}\$|\$\^[^$\s]+\$|\*\*[^*]+\*\*|\*[^*]+\*|\\textbf\{[^}]+\}|\\textit\{[^}]+\}|\\text\{[^}]+\})'
        parts = re.split(pattern, text)
        
        for part in parts:
            if not part:
                continue
                
            # Handle subscripts: $_2$ or $_{abc}$
            if re.match(r'\$_\{[^}]+\}\$', part):
                subscript_text = re.findall(r'\$_\{([^}]+)\}\$', part)[0]
                run = paragraph.add_run(subscript_text)
                run.font.subscript = True
            elif re.match(r'\$_[^$\s]+\$', part):
                subscript_text = re.findall(r'\$_([^$\s]+)\$', part)[0]
                run = paragraph.add_run(subscript_text)
                run.font.subscript = True
            
            # Handle superscripts: $^2$ or $^{abc}$
            elif re.match(r'\$\^\{[^}]+\}\$', part):
                superscript_text = re.findall(r'\$\^\{([^}]+)\}\$', part)[0]
                run = paragraph.add_run(superscript_text)
                run.font.superscript = True
            elif re.match(r'\$\^[^$\s]+\$', part):
                superscript_text = re.findall(r'\$\^([^$\s]+)\$', part)[0]
                run = paragraph.add_run(superscript_text)
                run.font.superscript = True
            
            # Handle bold text: **text** or \textbf{text}
            elif re.match(r'\*\*[^*]+\*\*', part):
                bold_text = re.findall(r'\*\*([^*]+)\*\*', part)[0]
                run = paragraph.add_run(bold_text)
                run.bold = True
            elif re.match(r'\\textbf\{[^}]+\}', part):
                bold_text = re.findall(r'\\textbf\{([^}]+)\}', part)[0]
                run = paragraph.add_run(bold_text)
                run.bold = True
            
            # Handle italic text: *text* or \textit{text}
            elif re.match(r'\*[^*]+\*', part):
                italic_text = re.findall(r'\*([^*]+)\*', part)[0]
                run = paragraph.add_run(italic_text)
                run.italic = True
            elif re.match(r'\\textit\{[^}]+\}', part):
                italic_text = re.findall(r'\\textit\{([^}]+)\}', part)[0]
                run = paragraph.add_run(italic_text)
                run.italic = True
            
            # Handle regular text in \text{} blocks
            elif re.match(r'\\text\{[^}]+\}', part):
                text_content = re.findall(r'\\text\{([^}]+)\}', part)[0]
                paragraph.add_run(text_content)
            
            # Regular text
            else:
                paragraph.add_run(part)
    
    def _convert_chemical_formulas(self, text: str) -> str:
        """
        Pre-process text to convert common chemical formula patterns to proper formatting.
        """
        # Convert common patterns like C2H6 to C$_2$H$_6$
        # But avoid converting numbers that are already part of measurements or temperatures
        text = re.sub(r'(?<!\d)([A-Z][a-z]?)(\d+)(?![°.\d])', r'\1$_\2$', text)
        
        # Handle complex chemical formulas with parentheses like Ca(OH)2
        text = re.sub(r'\(([^)]+)\)(\d+)', r'(\1)$_\2$', text)
        
        # Handle hydrates like CuSO4·5H2O
        text = re.sub(r'·(\d+)', r'·$_\1$', text)
        
        # Handle degree symbols and common scientific notation
        text = text.replace('°C', '°C')
        text = text.replace(' C)', ' °C)')
        text = text.replace('(C)', '(°C)')
        
        # Handle percentage symbols
        text = text.replace(' %', '%')
        
        # Fix common chemical names that shouldn't be formatted
        text = text.replace('CO$_2$', 'CO$_2$')  # Already correct
        text = text.replace('H$_2$O', 'H$_2$O')  # Already correct
        
        return text
    
    async def generate_final_documents(
        self, 
        request: BaseAIRequest,
        summary_text: str,
        proposal_text: str,
        smiles_string: str,
        structure_image_base64: str,
        molecule_name: str
    ) -> FinalDocumentsResponse:
        """Generate all final documents for the research proposal"""
        
        try:
            # Generate comprehensive proposal text
            full_proposal = await self._generate_full_proposal(
                request, summary_text, proposal_text, smiles_string, molecule_name
            )
            
            # Generate synthesis recipe
            recipe_base64 = await self._generate_recipe_docx(
                request, proposal_text, smiles_string, molecule_name
            )
            
            # Generate data recording template
            data_template_base64 = self._generate_data_template()
            
            # Generate Word document
            docx_base64 = self._generate_proposal_docx(
                full_proposal, structure_image_base64, molecule_name
            )
            
            return FinalDocumentsResponse(
                full_proposal_text=full_proposal,
                recipe_file_base64=recipe_base64,
                data_template_base64=data_template_base64,
                proposal_docx_base64=docx_base64
            )
            
        except Exception as e:
            raise HTTPException(
                status_code=500,
                detail=f"Error generating documents: {str(e)}"
            )
    
    async def _generate_full_proposal(
        self, 
        request: BaseAIRequest,
        summary_text: str,
        proposal_text: str,
        smiles_string: str,
        molecule_name: str
    ) -> str:
        """Generate comprehensive research proposal"""
        
        prompt = f"""You are a PhD-level chemist writing a detailed research proposal. Create a comprehensive, professional proposal that would be suitable for submission to a funding agency or academic institution.

STRUCTURE YOUR PROPOSAL WITH THESE SECTIONS:
1. Title
2. Abstract (150-200 words)
3. Introduction and Background (cite the literature findings)
4. Research Objectives and Hypotheses
5. Methodology and Experimental Design
6. Expected Outcomes and Impact
7. Timeline and Milestones
8. Budget Considerations
9. Risk Assessment and Mitigation
10. Conclusion

CONTEXT INFORMATION:
- Literature Summary: {summary_text}
- Core Research Idea: {proposal_text}
- Target Molecule: {molecule_name}
- Target Molecule SMILES: {smiles_string}

REQUIREMENTS:
- Write in a professional, academic tone
- Include specific experimental procedures
- Reference the literature findings in your introduction
- Make the proposal scientifically rigorous and feasible
- Include safety considerations
- Suggest analytical methods for characterization
- Discuss potential applications and broader impact
- Make it publication-quality writing

IMPORTANT FORMATTING GUIDELINES:
- Write chemical formulas clearly (e.g., H2SO4, CO2, CH4)
- Use standard scientific notation for temperatures (e.g., 25°C, 150°C)
- Include proper units throughout (mg, mL, °C, MPa)
- Avoid LaTeX formatting like $_2$ - just write CO2, not CO$_2$
- Use clear, readable text formatting

Generate a comprehensive research proposal that demonstrates deep understanding of the field and presents a compelling case for funding."""

        try:
            response = await ai_service.generate_response(request, prompt)
            return response
        except Exception as e:
            raise HTTPException(
                status_code=500,
                detail=f"Error generating full proposal: {str(e)}"
            )
    
    async def _generate_recipe_docx(
        self,
        request: BaseAIRequest,
        proposal_text: str,
        smiles_string: str,
        molecule_name: str
    ) -> str:
        """Generate synthesis recipe as Word document"""
        
        prompt = f"""Based on the research proposal and target molecule, create a detailed synthesis recipe. Structure your response as a comprehensive laboratory procedure with the following sections:

1. MATERIALS AND REAGENTS (list all chemicals with molecular weights, amounts, equivalents, CAS numbers, suppliers, and safety notes)
2. EQUIPMENT REQUIRED
3. SAFETY PRECAUTIONS
4. EXPERIMENTAL PROCEDURE (step-by-step synthesis)
5. PURIFICATION METHODS
6. CHARACTERIZATION TECHNIQUES
7. EXPECTED YIELD AND PURITY
8. TROUBLESHOOTING NOTES

Target Molecule: {molecule_name}
SMILES: {smiles_string}
Research Context: {proposal_text}

Provide realistic amounts for a small-scale laboratory synthesis (1-10 mmol scale). Make this a professional, detailed laboratory protocol that a graduate student could follow. Include specific temperatures, times, and conditions.

IMPORTANT FORMATTING GUIDELINES:
- Write chemical formulas clearly (e.g., H2SO4, NaOH, CaCl2)
- Use standard notation for temperatures (e.g., 80°C, room temperature)
- Include proper units (mg, mL, °C, hours)
- Be specific about concentrations and equivalents
- Avoid LaTeX formatting like $_2$ - just write H2O, not H$_2$O"""

        try:
            recipe_response = await ai_service.generate_response(request, prompt)
            
            # Create Word document
            doc = Document()
            
            # Set document style
            style = doc.styles['Normal']
            style.font.name = 'Times New Roman'
            
            # Add title
            title = doc.add_heading(f'Synthesis Recipe: {molecule_name}', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add SMILES information
            smiles_para = doc.add_paragraph()
            smiles_para.add_run('Target Molecule SMILES: ').bold = True
            smiles_para.add_run(smiles_string)
            smiles_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph()  # Add spacing
            
            # Process the AI response and format as document
            # First convert chemical formulas to proper format
            recipe_response = self._convert_chemical_formulas(recipe_response)
            
            sections = recipe_response.split('\n')
            current_section = None
            
            for line in sections:
                line = line.strip()
                if not line:
                    continue
                
                # Check for section headers
                if any(keyword in line.upper() for keyword in ['MATERIALS', 'EQUIPMENT', 'SAFETY', 'PROCEDURE', 'PURIFICATION', 'CHARACTERIZATION', 'YIELD', 'TROUBLESHOOTING']):
                    # Add as heading
                    if line.startswith('#'):
                        doc.add_heading(line[1:].strip(), level=2)
                    else:
                        doc.add_heading(line, level=2)
                    current_section = line.upper()
                elif line.startswith('##'):
                    doc.add_heading(line[2:].strip(), level=3)
                elif line.startswith('#'):
                    doc.add_heading(line[1:].strip(), level=2)
                elif line.startswith('-') or line.startswith('*'):
                    # Bullet point with proper formatting
                    para = doc.add_paragraph(style='List Bullet')
                    self._process_text_for_word(line[1:].strip(), para)
                elif len(line) >= 3 and line[0].isdigit() and line[1:3] in ['. ', ') ']:
                    # Numbered list with proper formatting
                    para = doc.add_paragraph(style='List Number')
                    self._process_text_for_word(line[2:].strip(), para)
                else:
                    # Regular paragraph with proper formatting
                    para = doc.add_paragraph()
                    self._process_text_for_word(line, para)
            
            # Add footer with generation info
            doc.add_page_break()
            footer_para = doc.add_paragraph()
            footer_para.add_run('Document generated by AI Chemistry Research Assistant').italic = True
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Save to bytes
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            return base64.b64encode(buffer.getvalue()).decode('utf-8')
            
        except Exception as e:
            raise HTTPException(
                status_code=500,
                detail=f"Error generating recipe Word document: {str(e)}"
            )
    
    def _generate_data_template(self) -> str:
        """Generate data recording template Excel file"""
        
        try:
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "Experimental Data"
            
            # Add headers for experimental tracking
            headers = [
                "Experiment ID", "Date", "Researcher", "Reaction Conditions",
                "Starting Material (mg)", "Reagent 1 (mg)", "Reagent 2 (mg)",
                "Solvent (mL)", "Temperature (°C)", "Reaction Time (h)",
                "Stirring Rate (rpm)", "Atmosphere", "Workup Method",
                "Crude Yield (mg)", "Purified Yield (mg)", "Yield (%)",
                "Purity (NMR/HPLC)", "Melting Point (°C)", "Spectral Data",
                "Notes", "Success Rating (1-5)"
            ]
            
            ws.append(headers)
            
            # Style headers
            for cell in ws[1]:
                cell.font = openpyxl.styles.Font(bold=True)
                cell.fill = openpyxl.styles.PatternFill(
                    start_color="4472C4", end_color="4472C4", fill_type="solid"
                )
                cell.font = openpyxl.styles.Font(color="FFFFFF", bold=True)
            
            # Add some example rows
            example_rows = [
                ["EXP-001", "2024-01-15", "Researcher Name", "Standard conditions", 
                 "100", "50", "25", "10", "80", "2", "500", "N2", "Aqueous workup",
                 "85", "78", "78%", ">95%", "145-147", "See attached", "Good reaction", "4"],
                ["EXP-002", "", "", "", "", "", "", "", "", "", "", "", "",
                 "", "", "", "", "", "", "", ""]
            ]
            
            for row in example_rows:
                ws.append(row)
            
            # Auto-adjust column widths
            for column in ws.columns:
                max_length = 0
                column_letter = column[0].column_letter
                for cell in column:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(str(cell.value))
                    except:
                        pass
                adjusted_width = min(max_length + 2, 30)
                ws.column_dimensions[column_letter].width = adjusted_width
            
            # Save to bytes
            buffer = io.BytesIO()
            wb.save(buffer)
            buffer.seek(0)
            
            return base64.b64encode(buffer.getvalue()).decode('utf-8')
            
        except Exception as e:
            raise HTTPException(
                status_code=500,
                detail=f"Error generating data template: {str(e)}"
            )
    
    def _generate_proposal_docx(
        self,
        full_proposal_text: str,
        structure_image_base64: str,
        molecule_name: str
    ) -> str:
        """Generate Word document with proposal and structure"""
        
        try:
            doc = Document()
            
            # Set document style
            style = doc.styles['Normal']
            style.font.name = 'Times New Roman'
            # Note: Cannot set font size this way in python-docx, skip this line
            
            # Add title
            title = doc.add_heading('AI-Generated Research Proposal', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add structure image if available
            if structure_image_base64:
                try:
                    # Decode base64 image
                    img_data = base64.b64decode(structure_image_base64)
                    image_stream = io.BytesIO(img_data)
                    
                    # Add centered paragraph for image
                    img_paragraph = doc.add_paragraph()
                    img_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = img_paragraph.runs[0] if img_paragraph.runs else img_paragraph.add_run()
                    run.add_picture(image_stream, width=Inches(4.0))
                    
                    # Add caption
                    caption = doc.add_paragraph(f"Figure 1: Target Molecular Structure - {molecule_name}")
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption.style = 'Caption'
                    
                except Exception as e:
                    print(f"Could not add image to document: {e}")
                    doc.add_paragraph(f"[Structure image could not be embedded - {molecule_name}]")
            
            doc.add_page_break()
            
            # Process proposal text
            # First convert chemical formulas to proper format
            full_proposal_text = self._convert_chemical_formulas(full_proposal_text)
            
            sections = full_proposal_text.split('\n')
            
            for line in sections:
                line = line.strip()
                if not line:
                    continue
                
                # Handle different heading levels
                if line.startswith('# '):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                elif line.startswith('#### '):
                    doc.add_heading(line[5:], level=4)
                elif line.upper() in ['ABSTRACT', 'INTRODUCTION', 'METHODOLOGY', 'CONCLUSION', 
                                    'RESULTS', 'DISCUSSION', 'REFERENCES', 'BACKGROUND',
                                    'OBJECTIVES', 'EXPERIMENTAL', 'TIMELINE', 'BUDGET']:
                    doc.add_heading(line.title(), level=2)
                elif ':' in line and len(line.split(':')[0]) < 50:  # Likely a section header
                    parts = line.split(':', 1)
                    if len(parts[0]) < 50:
                        doc.add_heading(parts[0].strip(), level=3)
                        if len(parts) > 1 and parts[1].strip():
                            # Use proper formatting for the content after colon
                            para = doc.add_paragraph()
                            self._process_text_for_word(parts[1].strip(), para)
                    else:
                        # Regular paragraph with proper formatting
                        para = doc.add_paragraph()
                        self._process_text_for_word(line, para)
                else:
                    # Regular paragraph with proper formatting
                    para = doc.add_paragraph()
                    self._process_text_for_word(line, para)
            
            # Save to bytes
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            return base64.b64encode(buffer.getvalue()).decode('utf-8')
            
        except Exception as e:
            raise HTTPException(
                status_code=500,
                detail=f"Error generating Word document: {str(e)}"
            )

# Global document service instance
document_service = DocumentService()